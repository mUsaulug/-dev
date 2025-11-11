#!/usr/bin/env python3
from docx import Document
import sys

def extract_docx_content_safe(filename):
    """Extract all content from a .docx file with error handling"""
    try:
        doc = Document(filename)
        content = {
            'filename': filename,
            'paragraphs': [],
            'tables': []
        }

        # Extract paragraphs with safe style handling
        for para in doc.paragraphs:
            if para.text.strip():
                try:
                    style_name = para.style.name if para.style else 'Normal'
                except:
                    style_name = 'Unknown'
                content['paragraphs'].append({
                    'text': para.text,
                    'style': style_name
                })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'index': table_idx,
                'data': table_data
            })

        return content
    except Exception as e:
        return {'error': str(e), 'filename': filename}

if __name__ == '__main__':
    filename = 'BYM_Technology_TUBITAK_Basvuru_FINAL.docx'

    print(f"{'='*80}")
    print(f"FILE: {filename}")
    print('='*80)
    content = extract_docx_content_safe(filename)

    if 'error' in content:
        print(f"ERROR: {content['error']}")
    else:
        print(f"\n--- PARAGRAPHS ({len(content['paragraphs'])}) ---")
        for i, para in enumerate(content['paragraphs'][:50]):  # First 50 paragraphs
            print(f"\n[{i+1}] ({para['style']}): {para['text'][:300]}")

        print(f"\n--- TABLES ({len(content['tables'])}) ---")
        for table in content['tables']:
            print(f"\nTable {table['index']+1} ({len(table['data'])} rows):")
            for row in table['data'][:3]:  # Show first 3 rows
                print(f"  {' | '.join([str(cell)[:60] for cell in row])}")
