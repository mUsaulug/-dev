#!/usr/bin/env python3
from docx import Document
import json
import sys

def extract_docx_content(filename):
    """Extract all content from a .docx file"""
    try:
        doc = Document(filename)
        content = {
            'filename': filename,
            'paragraphs': [],
            'tables': []
        }

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name
                })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'index': table_idx,
                'data': table_data
            })

        return content
    except Exception as e:
        return {'error': str(e), 'filename': filename}

if __name__ == '__main__':
    files = [
        'BYM_Technology_NIHAI_SUNUM.docx',
        'BYM_Technology_TUBITAK_Basvuru_FINAL.docx',
        'Kalp_Krizinde_Yapay_Zeka_Cozumleri_Final (1).docx'
    ]

    for filename in files:
        print(f"\n{'='*80}")
        print(f"FILE: {filename}")
        print('='*80)
        content = extract_docx_content(filename)

        if 'error' in content:
            print(f"ERROR: {content['error']}")
            continue

        print(f"\n--- PARAGRAPHS ({len(content['paragraphs'])}) ---")
        for i, para in enumerate(content['paragraphs']):
            print(f"\n[{i+1}] ({para['style']}): {para['text'][:200]}")

        print(f"\n--- TABLES ({len(content['tables'])}) ---")
        for table in content['tables']:
            print(f"\nTable {table['index']+1} ({len(table['data'])} rows):")
            for row in table['data'][:5]:  # Show first 5 rows
                print(f"  {' | '.join([str(cell)[:50] for cell in row])}")
