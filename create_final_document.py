#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BYM Technology - TÜBİTAK BİGG 1812 Nihai Sunum Dökümanı
Kalp Krizinde Yapay Zeka Çözümleri
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Sayfa sonu ekle"""
    doc.add_page_break()

def set_cell_border(cell, **kwargs):
    """Tablo hücresi kenarlık ayarları"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill):
    """Hücre arka plan rengi"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table_cell(cell, text, bold=False, size=20, color='000000', align='left'):
    """Formatlı tablo hücresi ekle"""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Yeni doküman oluştur
doc = Document()

# Sayfa marjları ayarla
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ═══════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ═══════════════════════════════════════════════════════════════════
title = doc.add_heading('KALP KRİZİNDE YAPAY ZEKA ÇÖZÜMLERİ', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(31, 71, 136)
title_run.font.size = Pt(36)

subtitle = doc.add_paragraph('Proaktif Kardiyovasküler Risk Yönetimi İçin AI-Destekli Karar Destek Sistemi')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
subtitle_run.italic = True

doc.add_paragraph()  # Boşluk
company = doc.add_paragraph('BYM TECHNOLOGY HEALTH CENTER')
company.alignment = WD_ALIGN_PARAGRAPH.CENTER
company_run = company.runs[0]
company_run.font.size = Pt(24)
company_run.font.color.rgb = RGBColor(46, 80, 144)
company_run.bold = True

program = doc.add_paragraph('TÜBİTAK BİGG 1812 Ürünleştirme ve Ticarileştirme Desteği')
program.alignment = WD_ALIGN_PARAGRAPH.CENTER
program_run = program.runs[0]
program_run.font.size = Pt(16)

date = doc.add_paragraph('Kasım 2025')
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date.runs[0]
date_run.font.size = Pt(14)
date_run.font.color.rgb = RGBColor(102, 102, 102)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 1. PROBLEMİN TANIMI
# ═══════════════════════════════════════════════════════════════════
heading1 = doc.add_heading('1. PROBLEMİN TANIMI', level=1)
heading1_run = heading1.runs[0]
heading1_run.font.color.rgb = RGBColor(46, 80, 144)

intro = doc.add_paragraph(
    'Kardiyovasküler hastalıklar küresel ölçekte 1 numaralı ölüm nedenidir (WHO, 2024). '
    'Mevcut sağlık sistemleri reaktif (kriz sonrası) müdahale yaparken, proaktif risk yönetimi '
    'eksikliği kritik kayıplara yol açmaktadır.'
)

# Problem İstatistikleri Tablosu
table = doc.add_table(rows=2, cols=3)
table.style = 'Light Grid Accent 1'

# Başlık satırı
header_cells = table.rows[0].cells
set_cell_shading(header_cells[0], '2E5090')
set_cell_shading(header_cells[1], '2E5090')
set_cell_shading(header_cells[2], '2E5090')
add_styled_table_cell(header_cells[0], 'Küresel Etki', True, 18, 'FFFFFF', 'center')
add_styled_table_cell(header_cells[1], 'Türkiye', True, 18, 'FFFFFF', 'center')
add_styled_table_cell(header_cells[2], 'Temel Sorunlar', True, 18, 'FFFFFF', 'center')

# Veri satırı
data_cells = table.rows[1].cells
# Küresel
para1 = data_cells[0].paragraphs[0]
para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1a = para1.add_run('17.9M\n')
run1a.bold = True
run1a.font.size = Pt(28)
run1a.font.color.rgb = RGBColor(211, 47, 47)
run1b = para1.add_run('yıllık ölüm\n')
run1b.font.size = Pt(16)
run1c = para1.add_run("Tüm ölümlerin %32'si (WHO, 2024)")
run1c.font.size = Pt(14)
run1c.italic = True

# Türkiye
para2 = data_cells[1].paragraphs[0]
para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2a = para2.add_run('489K\n')
run2a.bold = True
run2a.font.size = Pt(28)
run2a.font.color.rgb = RGBColor(211, 47, 47)
run2b = para2.add_run('yıllık ölüm\n')
run2b.font.size = Pt(16)
run2c = para2.add_run('1 numaralı ölüm sebebi (TÜİK, 2024)')
run2c.font.size = Pt(14)
run2c.italic = True

# Sorunlar
para3 = data_cells[2].paragraphs[0]
run3 = para3.add_run(
    '• Tanı hatası %10-12\n'
    '• Manuel süreçler\n'
    '• Veri entegrasyonu yok\n'
    '• Reaktif yaklaşım\n'
    '(Newman-Toker et al., 2024)'
)
run3.font.size = Pt(14)

# Sonuç bölümü
doc.add_heading('Sonuç', level=2)
conclusion = doc.add_paragraph(
    'Erken tanı ve proaktif risk yönetimi için yapay zeka destekli, çok katmanlı bir sistem '
    'kritik ihtiyaçtır. Kriz öncesi müdahale hayat kurtarır ve sağlık maliyetlerini düşürür.'
)
conclusion_run = conclusion.runs[0]
conclusion_run.bold = True

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 2. ÇÖZÜM: ÇOK KATMANLI AI SİSTEMİ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('2. ÇÖZÜM: ÇOK KATMANLI AI SİSTEMİ', level=1)

investor_pitch = doc.add_paragraph(
    'BYM Technology\'nin sistemi, üç veri katmanını (statik+vital+EKG) eş zamanlı analiz '
    'ederek proaktif risk değerlendirmesi yapar. SHAP/LIME ile açıklanabilir sonuçlar sunar. '
    'Sistem, kriz öncesi erken uyarı vererek hayat kurtarır ve hastane maliyetlerini %30-40 düşürür.'
)

# Sistem Modülleri Tablosu
table2 = doc.add_table(rows=5, cols=4)
table2.style = 'Light Grid Accent 1'

# Başlık
headers2 = table2.rows[0].cells
for cell in headers2:
    set_cell_shading(cell, '2E5090')
add_styled_table_cell(headers2[0], 'Modül', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers2[1], 'Teknoloji', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers2[2], 'Performans', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers2[3], 'Sonuç', True, 16, 'FFFFFF', 'center')

# Satır 1: Demografik
row1 = table2.rows[1].cells
set_cell_shading(row1[0], 'F2F2F2')
add_styled_table_cell(row1[0], 'Demografik Risk', True, 15)
add_styled_table_cell(row1[1], 'Bayes + AutoML', False, 14)
set_cell_shading(row1[2], 'E8F5E9')
add_styled_table_cell(row1[2], '%88+', True, 18, '4CAF50', 'center')
add_styled_table_cell(row1[3], 'Nedensel risk skoru', False, 14)

# Satır 2: EKG
row2 = table2.rows[2].cells
set_cell_shading(row2[0], 'F2F2F2')
add_styled_table_cell(row2[0], 'EKG Analizi', True, 15)
add_styled_table_cell(row2[1], 'CNN+LSTM+Attention', False, 13)
set_cell_shading(row2[2], 'E8F5E9')
add_styled_table_cell(row2[2], '%90+ (7 tip)', True, 16, '4CAF50', 'center')
add_styled_table_cell(row2[3], '<2sn gerçek zamanlı', False, 14)

# Satır 3: Koroner
row3 = table2.rows[3].cells
set_cell_shading(row3[0], 'F2F2F2')
add_styled_table_cell(row3[0], 'Koroner Görüntü', True, 15)
add_styled_table_cell(row3[1], '3D U-Net', False, 14)
set_cell_shading(row3[2], 'E8F5E9')
add_styled_table_cell(row3[2], '%87+ Dice', True, 16, '4CAF50', 'center')
add_styled_table_cell(row3[3], 'Oto tıkanıklık tespiti', False, 14)

# Satır 4: Açıklanabilir AI
row4 = table2.rows[4].cells
set_cell_shading(row4[0], 'FFF9C4')
set_cell_shading(row4[1], 'FFF9C4')
set_cell_shading(row4[2], 'FFF9C4')
set_cell_shading(row4[3], 'FFF9C4')
add_styled_table_cell(row4[0], 'Açıklanabilir AI', True, 15)
add_styled_table_cell(row4[1], 'SHAP + LIME', False, 14)
add_styled_table_cell(row4[2], 'Şeffaf', True, 16, '000000', 'center')
add_styled_table_cell(row4[3], 'Görsel gerekçelendirme', False, 13)

tech_note = doc.add_paragraph(
    'Teknik Altyapı: Python, TensorFlow, Keras | HL7/FHIR uyumlu | Modüler mimari'
)
tech_note_run = tech_note.runs[0]
tech_note_run.font.size = Pt(12)
tech_note_run.italic = True
tech_note_run.font.color.rgb = RGBColor(102, 102, 102)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 3. PROJE TAKVİMİ VE MİL TAŞLARI (YENİ GÖRSEL)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('3. PROJE TAKVİMİ VE MİL TAŞLARI', level=1)

# Timeline Tablosu
table3 = doc.add_table(rows=4, cols=4)
table3.style = 'Light Grid Accent 1'

# Başlık
headers3 = table3.rows[0].cells
for cell in headers3:
    set_cell_shading(cell, '2E5090')
add_styled_table_cell(headers3[0], 'Faz', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers3[1], 'Süre', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers3[2], 'Kilometre Taşı', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers3[3], 'Çıktılar', True, 16, 'FFFFFF', 'center')

# Faz 1
faz1 = table3.rows[1].cells
set_cell_shading(faz1[0], 'C5E1A5')
add_styled_table_cell(faz1[0], 'FAZ 1', True, 18, '33691E', 'center')
add_styled_table_cell(faz1[1], 'Ay 1-12', True, 15, '000000', 'center')
add_styled_table_cell(faz1[2], 'Sistem Geliştirme', True, 15)
p1 = faz1[3].paragraphs[0]
p1.add_run('✓ Algoritma tamamlama\n✓ Prototip cihaz\n✓ İlk validasyon').font.size = Pt(13)

# Faz 2
faz2 = table3.rows[2].cells
set_cell_shading(faz2[0], '90CAF9')
add_styled_table_cell(faz2[0], 'FAZ 2', True, 18, '0D47A1', 'center')
add_styled_table_cell(faz2[1], 'Ay 13-24', True, 15, '000000', 'center')
add_styled_table_cell(faz2[2], 'Pilot Uygulamalar', True, 15)
p2 = faz2[3].paragraphs[0]
p2.add_run('✓ 3 hastane pilotu\n✓ SaaS platformu\n✓ İlk müşteriler').font.size = Pt(13)

# Faz 3
faz3 = table3.rows[3].cells
set_cell_shading(faz3[0], 'FFB74D')
add_styled_table_cell(faz3[0], 'FAZ 3', True, 18, 'E65100', 'center')
add_styled_table_cell(faz3[1], 'Ay 25+', True, 15, '000000', 'center')
add_styled_table_cell(faz3[2], 'Ölçeklendirme', True, 15)
p3 = faz3[3].paragraphs[0]
p3.add_run('✓ CE/FDA süreçleri\n✓ Mobil uygulama\n✓ Uluslararası pazar').font.size = Pt(13)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 4. PAZAR ANALİZİ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('4. PAZAR ANALİZİ', level=1)

# Pazar Özeti Kutular
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'

cells4 = table4.rows[0].cells

# Küresel
set_cell_shading(cells4[0], 'E3F2FD')
p4a = cells4[0].paragraphs[0]
p4a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4a1 = p4a.add_run('KÜRESEL\n')
r4a1.bold = True
r4a1.font.size = Pt(16)
r4a1.font.color.rgb = RGBColor(25, 118, 210)
r4a2 = p4a.add_run('$32.1B\n')
r4a2.bold = True
r4a2.font.size = Pt(32)
r4a2.font.color.rgb = RGBColor(25, 118, 210)
r4a3 = p4a.add_run('2029 projeksiyon\n')
r4a3.font.size = Pt(13)
r4a4 = p4a.add_run('CAGR %14.6\n(McKinsey, 2024)')
r4a4.font.size = Pt(14)
r4a4.bold = True
r4a4.font.color.rgb = RGBColor(76, 175, 80)

# Türkiye
set_cell_shading(cells4[1], 'FFF3E0')
p4b = cells4[1].paragraphs[0]
p4b.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4b1 = p4b.add_run('TÜRKİYE\n')
r4b1.bold = True
r4b1.font.size = Pt(16)
r4b1.font.color.rgb = RGBColor(245, 124, 0)
r4b2 = p4b.add_run('$8.2B\n')
r4b2.bold = True
r4b2.font.size = Pt(32)
r4b2.font.color.rgb = RGBColor(245, 124, 0)
r4b3 = p4b.add_run('Sağlık Teknolojileri\n')
r4b3.font.size = Pt(13)
r4b4 = p4b.add_run('AI Payı %3.1\n(Statista, 2024)')
r4b4.font.size = Pt(14)
r4b4.bold = True
r4b4.font.color.rgb = RGBColor(255, 152, 0)

# Hedef
set_cell_shading(cells4[2], 'E8F5E9')
p4c = cells4[2].paragraphs[0]
p4c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4c1 = p4c.add_run('HEDEF\n')
r4c1.bold = True
r4c1.font.size = Pt(16)
r4c1.font.color.rgb = RGBColor(56, 142, 60)
r4c2 = p4c.add_run('372\n')
r4c2.bold = True
r4c2.font.size = Pt(32)
r4c2.font.color.rgb = RGBColor(56, 142, 60)
r4c3 = p4c.add_run('Özel Hastane\n')
r4c3.font.size = Pt(13)
r4c4 = p4c.add_run('2,840 Kardiyolog')
r4c4.font.size = Pt(14)
r4c4.bold = True

doc.add_heading('Müşteri Segmentleri', level=2)

# Müşteri Tablosu
table5 = doc.add_table(rows=4, cols=4)
table5.style = 'Light Grid Accent 1'

# Başlık
headers5 = table5.rows[0].cells
for cell in headers5:
    set_cell_shading(cell, '2E5090')
add_styled_table_cell(headers5[0], 'Paket', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers5[1], 'Yatak', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers5[2], 'Hedef', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers5[3], 'Fiyat/Ay', True, 16, 'FFFFFF', 'center')

# Temel
temel = table5.rows[1].cells
set_cell_shading(temel[0], 'F2F2F2')
add_styled_table_cell(temel[0], 'Temel', True, 15)
add_styled_table_cell(temel[1], '50-150', False, 14, '000000', 'center')
add_styled_table_cell(temel[2], '150', True, 18, '2E5090', 'center')
set_cell_shading(temel[3], 'E8F5E9')
add_styled_table_cell(temel[3], '₺10,000', True, 18, '000000', 'center')

# Profesyonel
prof = table5.rows[2].cells
set_cell_shading(prof[0], 'F2F2F2')
add_styled_table_cell(prof[0], 'Profesyonel', True, 15)
add_styled_table_cell(prof[1], '150-400', False, 14, '000000', 'center')
add_styled_table_cell(prof[2], '180', True, 18, '2E5090', 'center')
set_cell_shading(prof[3], 'E8F5E9')
add_styled_table_cell(prof[3], '₺20,000', True, 18, '000000', 'center')

# Kurumsal
kurum = table5.rows[3].cells
set_cell_shading(kurum[0], 'F2F2F2')
add_styled_table_cell(kurum[0], 'Kurumsal', True, 15)
add_styled_table_cell(kurum[1], '400+', False, 14, '000000', 'center')
add_styled_table_cell(kurum[2], '70', True, 18, '2E5090', 'center')
set_cell_shading(kurum[3], 'E8F5E9')
add_styled_table_cell(kurum[3], '₺35,000', True, 18, '000000', 'center')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 5. REKABET ÜSTÜNLÜĞÜ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('5. REKABET ÜSTÜNLÜĞÜ', level=1)

# Rakip Karşılaştırma
table6 = doc.add_table(rows=6, cols=5)
table6.style = 'Light Grid Accent 1'

# Başlık
headers6 = table6.rows[0].cells
set_cell_shading(headers6[0], '2E5090')
add_styled_table_cell(headers6[0], 'Özellik', True, 14, 'FFFFFF', 'center')
set_cell_shading(headers6[1], '4CAF50')
add_styled_table_cell(headers6[1], 'BYM Tech', True, 14, 'FFFFFF', 'center')
set_cell_shading(headers6[2], '2E5090')
add_styled_table_cell(headers6[2], 'CureMD', True, 14, 'FFFFFF', 'center')
set_cell_shading(headers6[3], '2E5090')
add_styled_table_cell(headers6[3], 'Merge', True, 14, 'FFFFFF', 'center')
set_cell_shading(headers6[4], '2E5090')
add_styled_table_cell(headers6[4], 'Cardisio', True, 14, 'FFFFFF', 'center')

# Kriz Öncesi
kr1 = table6.rows[1].cells
set_cell_shading(kr1[0], 'F2F2F2')
add_styled_table_cell(kr1[0], 'Kriz Öncesi', True, 13)
set_cell_shading(kr1[1], 'E8F5E9')
add_styled_table_cell(kr1[1], '✓ VAR', True, 14, '2E7D32', 'center')
set_cell_shading(kr1[2], 'FFEBEE')
add_styled_table_cell(kr1[2], '✗ YOK', False, 13, '000000', 'center')
set_cell_shading(kr1[3], 'FFEBEE')
add_styled_table_cell(kr1[3], '✗ YOK', False, 13, '000000', 'center')
set_cell_shading(kr1[4], 'FFF9C4')
add_styled_table_cell(kr1[4], 'Kısıtlı', False, 13, '000000', 'center')

# Çok Katman
kr2 = table6.rows[2].cells
set_cell_shading(kr2[0], 'F2F2F2')
add_styled_table_cell(kr2[0], 'Çok Katman', True, 13)
set_cell_shading(kr2[1], 'E8F5E9')
add_styled_table_cell(kr2[1], '3 Katman', True, 14, '2E7D32', 'center')
set_cell_shading(kr2[2], 'FFEBEE')
add_styled_table_cell(kr2[2], 'EHR', False, 13, '000000', 'center')
set_cell_shading(kr2[3], 'FFF9C4')
add_styled_table_cell(kr2[3], '2 Katman', False, 12, '000000', 'center')
set_cell_shading(kr2[4], 'FFEBEE')
add_styled_table_cell(kr2[4], 'Tek', False, 13, '000000', 'center')

# Açıklanabilir
kr3 = table6.rows[3].cells
set_cell_shading(kr3[0], 'F2F2F2')
add_styled_table_cell(kr3[0], 'Açıklanabilir', True, 13)
set_cell_shading(kr3[1], 'E8F5E9')
add_styled_table_cell(kr3[1], 'SHAP+LIME', True, 12, '2E7D32', 'center')
set_cell_shading(kr3[2], 'FFEBEE')
add_styled_table_cell(kr3[2], 'Yok', False, 13, '000000', 'center')
set_cell_shading(kr3[3], 'FFEBEE')
add_styled_table_cell(kr3[3], 'Yok', False, 13, '000000', 'center')
set_cell_shading(kr3[4], 'FFEBEE')
add_styled_table_cell(kr3[4], 'Yok', False, 13, '000000', 'center')

# Yerelleşme
kr4 = table6.rows[4].cells
set_cell_shading(kr4[0], 'F2F2F2')
add_styled_table_cell(kr4[0], 'Yerelleşme', True, 13)
set_cell_shading(kr4[1], 'E8F5E9')
add_styled_table_cell(kr4[1], 'Türkiye', True, 14, '2E7D32', 'center')
set_cell_shading(kr4[2], 'FFEBEE')
add_styled_table_cell(kr4[2], 'Global', False, 13, '000000', 'center')
set_cell_shading(kr4[3], 'FFEBEE')
add_styled_table_cell(kr4[3], 'Global', False, 13, '000000', 'center')
set_cell_shading(kr4[4], 'FFEBEE')
add_styled_table_cell(kr4[4], 'Global', False, 13, '000000', 'center')

# Maliyet
kr5 = table6.rows[5].cells
set_cell_shading(kr5[0], 'F2F2F2')
add_styled_table_cell(kr5[0], 'Maliyet', True, 13)
set_cell_shading(kr5[1], 'E8F5E9')
add_styled_table_cell(kr5[1], '%60 Düşük', True, 14, '2E7D32', 'center')
set_cell_shading(kr5[2], 'FFEBEE')
add_styled_table_cell(kr5[2], 'Yüksek', False, 13, '000000', 'center')
set_cell_shading(kr5[3], 'FFEBEE')
add_styled_table_cell(kr5[3], 'Yüksek', False, 13, '000000', 'center')
set_cell_shading(kr5[4], 'FFF9C4')
add_styled_table_cell(kr5[4], 'Orta', False, 13, '000000', 'center')

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 6. GELİR MODELİ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('6. GELİR MODELİ (5 YIL)', level=1)

# Gelir Akışları
table7 = doc.add_table(rows=6, cols=3)
table7.style = 'Light Grid Accent 1'

# Başlık
headers7 = table7.rows[0].cells
for cell in headers7:
    set_cell_shading(cell, '2E5090')
add_styled_table_cell(headers7[0], 'Gelir Kaynağı', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers7[1], 'Birim Fiyat', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers7[2], 'Toplam Gelir', True, 16, 'FFFFFF', 'center')

# SaaS
saas = table7.rows[1].cells
set_cell_shading(saas[0], 'F2F2F2')
add_styled_table_cell(saas[0], 'SaaS Abonelik', True, 16)
add_styled_table_cell(saas[1], '₺10K-35K/ay', False, 14, '000000', 'center')
set_cell_shading(saas[2], 'E8F5E9')
add_styled_table_cell(saas[2], '₺410M', True, 20, '2E7D32', 'center')

# Entegrasyon
ent = table7.rows[2].cells
set_cell_shading(ent[0], 'F2F2F2')
add_styled_table_cell(ent[0], 'Entegrasyon', True, 16)
add_styled_table_cell(ent[1], '₺120K-350K', False, 14, '000000', 'center')
set_cell_shading(ent[2], 'E8F5E9')
add_styled_table_cell(ent[2], '₺30M', True, 20, '2E7D32', 'center')

# Eğitim
egit = table7.rows[3].cells
set_cell_shading(egit[0], 'F2F2F2')
add_styled_table_cell(egit[0], 'Eğitim', True, 16)
add_styled_table_cell(egit[1], '₺6.5K/kişi', False, 14, '000000', 'center')
set_cell_shading(egit[2], 'E8F5E9')
add_styled_table_cell(egit[2], '₺8M', True, 20, '2E7D32', 'center')

# Veri
veri = table7.rows[4].cells
set_cell_shading(veri[0], 'F2F2F2')
add_styled_table_cell(veri[0], 'Veri Analizi', True, 16)
add_styled_table_cell(veri[1], '₺12K-20K/ay', False, 14, '000000', 'center')
set_cell_shading(veri[2], 'E8F5E9')
add_styled_table_cell(veri[2], '₺12M', True, 20, '2E7D32', 'center')

# Toplam
total = table7.rows[5].cells
set_cell_shading(total[0], '1B5E20')
add_styled_table_cell(total[0], 'TOPLAM', True, 18, 'FFFFFF', 'right')
set_cell_shading(total[1], '1B5E20')
add_styled_table_cell(total[1], '-', False, 14, 'FFFFFF', 'center')
set_cell_shading(total[2], '1B5E20')
add_styled_table_cell(total[2], '₺460M', True, 24, 'FFFFFF', 'center')

financial = doc.add_paragraph('Birim maliyet: ₺4,800 | Brüt kar: %108 | Başa baş: Yıl 3 | Hedef: 400 hastane')
financial_run = financial.runs[0]
financial_run.bold = True
financial_run.italic = True
financial_run.font.size = Pt(14)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 7. KLİNİK VALİDASYON
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('7. KLİNİK VALİDASYON VE DANIŞMANLIK', level=1)

# Danışmanlar
table8 = doc.add_table(rows=1, cols=2)
table8.style = 'Light Grid Accent 1'

# Dr. Mahsa
mahsa = table8.rows[0].cells[0]
set_cell_shading(mahsa, 'E1F5FE')
pm = mahsa.paragraphs[0]
pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
pm.add_run('👨‍⚕️\n\n').font.size = Pt(32)
rm1 = pm.add_run('Uzm. Dr. Mahsa Sedaghatıhagh\n')
rm1.bold = True
rm1.font.size = Pt(16)
rm2 = pm.add_run('Özel Doğa Hospital\n\n')
rm2.italic = True
rm2.font.size = Pt(13)
rm3 = pm.add_run('Klinik validasyon, hasta profilleri,\nkarar destek mekanizmaları')
rm3.font.size = Pt(12)

# Dr. Asil
asil = table8.rows[0].cells[1]
set_cell_shading(asil, 'F3E5F5')
pa = asil.paragraphs[0]
pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
pa.add_run('👨‍⚕️\n\n').font.size = Pt(32)
ra1 = pa.add_run('Uzm. Dr. Asil İşçi\n')
ra1.bold = True
ra1.font.size = Pt(16)
ra2 = pa.add_run('İSTÜN Kolan Hospital\n\n')
ra2.italic = True
ra2.font.size = Pt(13)
ra3 = pa.add_run('Klinik uygulanabilirlik,\ngirişimsel kardiyoloji,\n10K+ anjiografi deneyimi')
ra3.font.size = Pt(12)

validation = doc.add_paragraph('✓ 500 hasta verisi | ✓ ROC-AUC >0.90 | ✓ 3 pilot hastane | ✓ 60 gün gerçek test')
validation_run = validation.runs[0]
validation_run.bold = True
validation_run.font.size = Pt(14)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 8. KURUMSAL İŞBİRLİKLERİ
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('8. KURUMSAL İŞBİRLİKLERİ VE DESTEKLER', level=1)

# İşbirlikleri
table9 = doc.add_table(rows=4, cols=3)
table9.style = 'Light Grid Accent 1'

# Başlık
headers9 = table9.rows[0].cells
for cell in headers9:
    set_cell_shading(cell, '2E5090')
add_styled_table_cell(headers9[0], 'Kurum', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers9[1], 'Kategori', True, 16, 'FFFFFF', 'center')
add_styled_table_cell(headers9[2], 'Destek Kapsamı', True, 16, 'FFFFFF', 'center')

# İTÜ
itu = table9.rows[1].cells
set_cell_shading(itu[0], 'E3F2FD')
add_styled_table_cell(itu[0], 'İstanbul Topkapı Üniversitesi', True, 14, '000000', 'center')
add_styled_table_cell(itu[1], 'Akademik', True, 15, '1976D2', 'center')
pi = itu[2].paragraphs[0]
pi.add_run('• Ar-Ge desteği\n• EKG modeli araştırmaları\n• Akademik yayın').font.size = Pt(13)

# Doğa
doga = table9.rows[2].cells
set_cell_shading(doga[0], 'FCE4EC')
add_styled_table_cell(doga[0], 'Özel Doğa Hospital', True, 14, '000000', 'center')
add_styled_table_cell(doga[1], 'Klinik', True, 15, 'C2185B', 'center')
pd = doga[2].paragraphs[0]
pd.add_run('• Klinik validasyon\n• Hasta verisi erişimi\n• Pilot uygulama').font.size = Pt(13)

# Kolan
kolan = table9.rows[3].cells
set_cell_shading(kolan[0], 'F3E5F5')
add_styled_table_cell(kolan[0], 'İSTÜN Kolan Hospital', True, 14, '000000', 'center')
add_styled_table_cell(kolan[1], 'Klinik', True, 15, '7B1FA2', 'center')
pk = kolan[2].paragraphs[0]
pk.add_run('• Girişimsel validasyon\n• Gerçek ortam testi\n• Uzman feedback').font.size = Pt(13)

support = doc.add_paragraph('✓ İyi niyet mektupları ekli | ✓ MoU imza sürecinde | ✓ Pilot uygulamalar hazır')
support_run = support.runs[0]
support_run.bold = True
support_run.font.size = Pt(14)
support_run.font.color.rgb = RGBColor(76, 175, 80)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# 9. KAYNAKÇA
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('9. KAYNAKÇA', level=1)

references = [
    'WHO (2024). Cardiovascular diseases fact sheet. World Health Organization.',
    'TÜİK (2024). Ölüm nedeni istatistikleri. Türkiye İstatistik Kurumu.',
    'Newman-Toker, D.E. et al. (2024). Diagnostic errors in cardiovascular emergencies. JAMA Internal Medicine.',
    'Grand View Research (2024). Artificial Intelligence in Healthcare Market Size & Trends.',
    'Fortune Business Insights (2024). Healthcare AI Market Growth Analysis.',
    'McKinsey & Company (2024). The State of AI in Healthcare.',
    'Statista (2024). Healthcare Technology Market in Turkey.',
    'IMARC Group (2024). Turkey Healthcare Market Report.',
    'MarketsandMarkets (2024). Medical Imaging AI Market Forecast.',
    'T.C. Sağlık Bakanlığı (2024). Sağlık İstatistikleri Yıllığı.'
]

for ref in references:
    p = doc.add_paragraph(ref, style='List Bullet')
    p.runs[0].font.size = Pt(12)

add_page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# ARKA KAPAK - İLETİŞİM
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()  # Boşluklar
doc.add_paragraph()
doc.add_paragraph()

contact1 = doc.add_paragraph('BYM TECHNOLOGY')
contact1.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact1_run = contact1.runs[0]
contact1_run.font.size = Pt(32)
contact1_run.bold = True
contact1_run.font.color.rgb = RGBColor(46, 80, 144)

contact2 = doc.add_paragraph('HEALTH CENTER')
contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact2_run = contact2.runs[0]
contact2_run.font.size = Pt(28)
contact2_run.bold = True
contact2_run.font.color.rgb = RGBColor(46, 80, 144)

doc.add_paragraph()
doc.add_paragraph()

email = doc.add_paragraph('📧 bymtechnology3@gmail.com')
email.alignment = WD_ALIGN_PARAGRAPH.CENTER
email_run = email.runs[0]
email_run.font.size = Pt(18)
email_run.bold = True

loc1 = doc.add_paragraph('📍 İstanbul Topkapı Üniversitesi')
loc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
loc1_run = loc1.runs[0]
loc1_run.font.size = Pt(14)

loc2 = doc.add_paragraph('Teknoloji Geliştirme Bölgesi')
loc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
loc2_run = loc2.runs[0]
loc2_run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

prog1 = doc.add_paragraph('TÜBİTAK BİGG 1812 Programı')
prog1.alignment = WD_ALIGN_PARAGRAPH.CENTER
prog1_run = prog1.runs[0]
prog1_run.font.size = Pt(16)
prog1_run.bold = True

prog2 = doc.add_paragraph('Ürünleştirme ve Ticarileştirme Desteği')
prog2.alignment = WD_ALIGN_PARAGRAPH.CENTER
prog2_run = prog2.runs[0]
prog2_run.font.size = Pt(14)

doc.add_paragraph()

final_date = doc.add_paragraph('Kasım 2025')
final_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
final_date_run = final_date.runs[0]
final_date_run.font.size = Pt(14)
final_date_run.font.color.rgb = RGBColor(102, 102, 102)

# Kaydet
output_path = '/home/user/-dev/BYM_Technology_NIHAI_SUNUM.docx'
doc.save(output_path)
print(f'✅ NİHAİ DOKÜMAN OLUŞTURULDU: {output_path}')
